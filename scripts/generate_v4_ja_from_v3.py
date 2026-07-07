#!/usr/bin/env python3
"""
Generate v4 JA manuscript from JA v3 base.
Same strategy as EN: read v3, make targeted modifications, insert new sections.
"""

import os
import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.join(SCRIPT_DIR, '..')
FIG_DIR = os.path.join(ROOT_DIR, 'figures')
OUT_DIR = os.path.join(ROOT_DIR, 'manuscript')

with open(os.path.join(FIG_DIR, 'revision_results.json'), 'r') as f:
    RESULTS = json.load(f)
with open(os.path.join(FIG_DIR, 'per_polity_ks.json'), 'r') as f:
    KS_PVALUES = json.load(f)


def find_ja_v3():
    paths = [
        os.path.join(OUT_DIR, 'hassc_manuscript_ja_v3.docx'),
        '/home/ubuntu/attachments/6300c229-7ff0-45a2-8d01-7d5c283237c1/hassc_manuscript_ja.docx',
    ]
    for p in paths:
        if os.path.exists(p):
            return p
    sys.exit("ERROR: Cannot find JA v3")


def add_paragraph_after(doc, ref_paragraph, text):
    new_p = OxmlElement('w:p')
    ref_paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, doc)
    if text:
        run = new_para.add_run(text)
        run.font.size = Pt(11)
    return new_para


def add_heading_after(doc, ref_paragraph, text, level=2):
    new_p = OxmlElement('w:p')
    ref_paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, doc)
    new_para.style = doc.styles[f'Heading {level}']
    run = new_para.add_run(text)
    return new_para


def replace_paragraph_text_fully(paragraph, new_text):
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        run = paragraph.add_run(new_text)
        run.font.size = Pt(11)


def generate_ja_v4():
    v3_path = find_ja_v3()
    print(f"Reading JA v3 from: {v3_path}")
    doc = Document(v3_path)
    print(f"JA V3 has {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    # Build paragraph map
    para_map = {}
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
        if txt.startswith('ワイブル分布は、機械部品') and i < 15:
            para_map['abstract'] = i
        elif 'Palgrave Communications誌に掲載された研究' in txt:
            para_map['intro_saleh'] = i
        elif txt.startswith('各政体は支配的な継承メカニズム'):
            para_map['classification'] = i
        elif txt.startswith('ワイブル分布は分析した31政体の大多数'):
            para_map['results_main1'] = i
        elif txt.startswith('三つの異なる安定性体制'):
            para_map['results_main2'] = i
        elif txt.startswith('歴史的検証：暴力的継承') and p.style.name.startswith('Heading'):
            para_map['results_histval_heading'] = i
        elif txt.startswith('ワイブル形状パラメータkの解釈を検証'):
            para_map['results_histval_body'] = i
        elif txt.startswith('図6. 歴史的検証'):
            para_map['fig6_caption'] = i
        elif txt.startswith('サブ解析1：選挙制') and p.style.name.startswith('Heading'):
            para_map['results_elective_heading'] = i
        elif txt.startswith('31政体を選挙制'):
            para_map['results_elective_body'] = i
        elif txt.startswith('継承メカニズム別のサブ解析'):
            para_map['discussion_elective'] = i
        elif txt.startswith('歴史的検証分析はkを安定性指標'):
            para_map['discussion_histval'] = i
        elif txt == '限界' and p.style.name.startswith('Heading'):
            para_map['limitations_heading'] = i
        elif txt == '考察' and p.style.name == 'Heading 1':
            para_map['discussion_heading'] = i
        elif txt.startswith('第五に、年代への在位'):
            para_map['future_work'] = i
        elif txt.startswith('本研究は、信頼性工学の基幹ツール'):
            para_map['conclusion'] = i
        elif txt == 'ブートストラップ信頼区間' and p.style.name.startswith('Heading'):
            para_map['bootstrap_heading'] = i
        elif txt == '統計的比較' and p.style.name.startswith('Heading'):
            para_map['stat_heading'] = i
        elif txt.startswith('感度分析：幼帝の除外') and p.style.name.startswith('Heading'):
            para_map['sensitivity_heading'] = i

    print(f"Mapped {len(para_map)} key paragraphs")
    for k, v in sorted(para_map.items(), key=lambda x: x[1]):
        print(f"  {k}: para {v}")

    ms = RESULTS['model_selection']
    papacy = RESULTS['papacy']
    saleh = RESULTS['roman_saleh']
    usa = RESULTS['usa_sensitivity']
    ss = RESULTS['sample_size']
    n_total = len(ms)
    w_best_aic = sum(1 for r in ms.values() if r['best_aic'] == 'Weibull')
    w_comp = sum(1 for r in ms.values() if r.get('w_daic') is not None and r['w_daic'] < 2)
    n_pass = sum(1 for v in KS_PVALUES.values() if v['ks_p'] > 0.05)

    # =========================================================================
    # MODIFICATION: Classification - reclassify USA
    # =========================================================================
    if 'classification' in para_map:
        p = doc.paragraphs[para_map['classification']]
        old_text = p.text
        new_text = old_text.replace(
            '選挙制政体（N = 8）',
            '選挙制政体（N = 7）'
        )
        if 'アメリカ合衆国大統領' in new_text:
            new_text = new_text.replace(
                'アメリカ合衆国大統領',
                'アメリカ合衆国大統領（※任期制限付き民主制として別分類、後述）'
            )
        replace_paragraph_text_fully(p, new_text)
        print("R2.2: USA reclassified")

    # =========================================================================
    # MODIFICATION: Results main1 - actual p-values
    # =========================================================================
    if 'results_main1' in para_map:
        p = doc.paragraphs[para_map['results_main1']]
        old_text = p.text
        new_text = old_text.replace(
            'KS検定 p > 0.05',
            f'KS検定 p > 0.05; {n_pass}/{n_total}政体で適合'
        )
        replace_paragraph_text_fully(p, new_text)
        print("Fixed p-values in results")

    # =========================================================================
    # MODIFICATION: Results main2 - add CI crossing note
    # =========================================================================
    if 'results_main2' in para_map:
        p = doc.paragraphs[para_map['results_main2']]
        old_text = p.text
        new_text = old_text + (
            f' なお、31政体中{ss["n_crossing"]}政体の95%信頼区間がk = 1を跨いでおり、'
            'これらの政体の安定性体制分類は統計的に決定的ではない。'
            'また、N < 15の4政体（アステカ帝国 N = 9、タイ・チャクリー朝 N = 10、'
            '中国・清 N = 11、インカ帝国 N = 12）は信頼区間が広く、定性的解釈は慎重を要する。'
        )
        replace_paragraph_text_fully(p, new_text)
        print("R2.3: Sample size threshold added")

    # =========================================================================
    # MODIFICATION: Elective body - USA sensitivity
    # =========================================================================
    if 'results_elective_body' in para_map:
        p = doc.paragraphs[para_map['results_elective_body']]
        old_text = p.text
        new_text = old_text.replace(
            '選挙制（8政体、プール446件）',
            '選挙制（7政体、プール404件、アメリカ合衆国大統領を除く）'
        )
        new_text += (
            f' 感度分析として、アメリカ合衆国を含めた場合のプール選挙制k = {usa["with"]:.3f}、'
            f'除外した場合のk = {usa["without"]:.3f}（差 = {usa["delta"]:.4f}）であり、'
            'アメリカ合衆国の任期制限による影響はプール推定値に対して最小限であることが確認された。'
        )
        replace_paragraph_text_fully(p, new_text)
        print("R2.2: Elective body updated")

    # =========================================================================
    # MODIFICATION: Reframe "歴史的検証" → "内部整合性"
    # =========================================================================
    if 'results_histval_heading' in para_map:
        p = doc.paragraphs[para_map['results_histval_heading']]
        replace_paragraph_text_fully(p, '内部整合性の検証：暴力的継承・内戦との相関')
        print("R2.minor-a: Results heading reframed")

    if 'results_histval_body' in para_map:
        p = doc.paragraphs[para_map['results_histval_body']]
        old_text = p.text
        new_text = old_text.replace('検証するため', '評価するため')
        new_text += (
            ' ワイブル形状パラメータと暴力的継承率はいずれも同一の歴史的記録から導出されるため'
            '（在位期間の境界を決定するイベントが継承の暴力性の分類にも影響する）、'
            'この分析は独立した外部検証ではなく内部整合性の実証として位置づけられる。'
        )
        replace_paragraph_text_fully(p, new_text)
        print("R2.minor-a: Results body reframed")

    if 'fig6_caption' in para_map:
        p = doc.paragraphs[para_map['fig6_caption']]
        old_text = p.text
        new_text = old_text.replace('歴史的検証', '内部整合性の検証')
        replace_paragraph_text_fully(p, new_text)
        print("R2.minor-a: Fig 6 caption reframed")

    if 'discussion_histval' in para_map:
        p = doc.paragraphs[para_map['discussion_histval']]
        old_text = p.text
        new_text = old_text.replace(
            '歴史的検証分析はkを安定性指標として解釈することに強い実証的支持を提供する',
            '内部整合性分析はkが実質的な歴史的意味を持つことを示す'
        )
        replace_paragraph_text_fully(p, new_text)
        print("R2.minor-a: Discussion reframed")

    if 'conclusion' in para_map:
        p = doc.paragraphs[para_map['conclusion']]
        old_text = p.text
        new_text = old_text.replace('歴史的検証', '内部整合性分析')
        replace_paragraph_text_fully(p, new_text)
        print("R2.minor-a: Conclusion reframed")

    # =========================================================================
    # MODIFICATION: Discussion elective - add USA discussion
    # =========================================================================
    if 'discussion_elective' in para_map:
        p = doc.paragraphs[para_map['discussion_elective']]
        old_text = p.text
        new_text = old_text + (
            ' アメリカ合衆国大統領（k = 2.721）はデータセット中で独自の位置を占める。'
            '憲法上の任期制限（当初は事実上、1951年の修正第22条で正式化）が在位期間を機械的に制約し、'
            '有機的な政治的ダイナミクスではなく制度設計によりkを押し上げている。'
            'このため、アメリカ合衆国大統領は表1に比較目的で報告するが、'
            'プール選挙制分析からは除外し、「任期制限付き民主制」として別カテゴリで扱う。'
        )
        replace_paragraph_text_fully(p, new_text)
        print("R2.2: Discussion USA updated")

    # =========================================================================
    # INSERT: Model selection in Methods (before stat_heading)
    # =========================================================================
    if 'stat_heading' in para_map:
        ref = doc.paragraphs[para_map['stat_heading'] - 1]

        ms_body_text = (
            'ワイブル分布が単に許容可能な当てはまりを提供するだけでなく、最良の当てはまりを提供するかを'
            '評価するため、正の右裾データに対する4つの候補分布（ワイブル、対数正規、ガンマ、指数）の'
            '体系的なモデル選択を実施した。すべての分布は位置パラメータをゼロに固定してMLEで当てはめた。'
            'モデル比較には赤池情報量規準（AIC）およびベイズ情報量規準（BIC）を用いた。'
            'ΔAIC < 2は本質的に同等の当てはまりを示し、ΔAIC 2-10は中程度の証拠、'
            'ΔAIC > 10は強い証拠を示す（Burnham and Anderson, 2002）。'
            '指数分布（1パラメータ）は無記憶な継承（k = 1）に対応するネスト帰無モデルとして機能する。'
        )
        ms_body = add_paragraph_after(doc, ref, None)
        replace_paragraph_text_fully(ms_body, ms_body_text)

        add_heading_after(doc, ref, 'モデル選択', level=2)

        # Also insert sample size methods
        ss_body_text = (
            'サンプルサイズがN < 15の政体（アステカ帝国 N = 9、タイ・チャクリー朝 N = 10、'
            '中国・清 N = 11、インカ帝国 N = 12）は広い信頼区間を生じ、信頼性の高い定性的解釈を'
            '排除する。これらの政体は表1に完全性のために含めるが、フラグを付け、'
            '定性的な体制分類からは除外する。また、kの95%信頼区間がk = 1の閾値を跨ぐ政体は、'
            '安定性体制分類が不確定であると注記する。'
        )
        ss_body = add_paragraph_after(doc, ms_body, None)
        replace_paragraph_text_fully(ss_body, ss_body_text)
        add_heading_after(doc, ms_body, '最小サンプルサイズに関する考慮', level=2)
        print("Inserted: Model selection + sample size methods")

    # =========================================================================
    # INSERT: Model selection results (before internal consistency heading)
    # =========================================================================
    results_insert_ref = None
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt.startswith('内部整合性') and p.style.name.startswith('Heading'):
            results_insert_ref = doc.paragraphs[i - 1]
            break

    if results_insert_ref is not None:
        ln_best = sum(1 for r in ms.values() if r['best_aic'] == 'Log-normal')
        g_best = sum(1 for r in ms.values() if r['best_aic'] == 'Gamma')
        e_best = sum(1 for r in ms.values() if r['best_aic'] == 'Exponential')

        disfavoured = [(n, r) for n, r in ms.items()
                       if r.get('w_daic') is not None and r['w_daic'] > 4]
        disfavoured_text = '、'.join(
            f'{n}（ΔAIC = {r["w_daic"]:.1f}、{r["best_aic"]}が優位）'
            for n, r in sorted(disfavoured, key=lambda x: -x[1]['w_daic'])
        )

        ms_results_text = (
            f'31政体すべてに対してワイブル、対数正規、ガンマ、指数分布の体系的モデル選択を実施した。'
            f'AICにより、ワイブル分布は{w_best_aic}/{n_total}政体で最良モデルであり、'
            f'対数正規が{ln_best}、ガンマが{g_best}、指数が{e_best}政体で最良であった。'
            f'重要なことに、ワイブルは{w_comp}/{n_total}政体でΔAIC < 2以内であり、'
            f'データセットの大多数で統計的に競争力のある当てはまりを提供することが確認された。'
            f'ワイブルが明確に劣位となった政体（ΔAIC > 4）は{disfavoured_text}であった。'
            f'日本天皇については在位期間の重い右裾（50年を超える天皇が複数存在）が対数正規分布の'
            f'より重い裾でよく捕捉される。アメリカ合衆国大統領については任期制限による離散的構造が'
            f'対数正規でより自然に記述される。これらの結果はワイブル分布が文化横断的比較のための'
            f'妥当な分析的枠組みを提供することを確認しつつ、普遍的に最良の当てはまりモデルではない'
            f'ことを認める。形状パラメータkの解釈上の利点（バスタブ曲線の枠組みを介したハザード'
            f'動態への直接的対応）は、代替分布がわずかに良い統計的当てはまりを達成する場合でも、'
            f'主要指標としてのkの使用を正当化する。図9にすべての政体と分布のΔAIC・ΔBIC値を示す。'
        )

        ms_fig_caption = add_paragraph_after(doc, results_insert_ref, None)
        replace_paragraph_text_fully(ms_fig_caption,
            '図9. モデル選択比較。各政体について最良モデルに対するΔAIC（左）とΔBIC（右）。'
            '塗り潰し丸は最良モデル、白抜き四角は競争的代替モデル（ΔAIC < 2）を示す。')
        ms_fig_caption.runs[0].italic = True

        fig_path = os.path.join(FIG_DIR, 'fig_model_selection.png')
        if os.path.exists(fig_path):
            img_p = OxmlElement('w:p')
            ms_fig_caption._element.addprevious(img_p)
            from docx.text.paragraph import Paragraph
            img_para = Paragraph(img_p, doc)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(fig_path, width=Inches(6.0))

        # S1 reference
        s1_ref = add_paragraph_after(doc, ms_fig_caption, None)
        replace_paragraph_text_fully(s1_ref,
            '補足図S1に全31政体のカプラン・マイヤー経験的生存曲線と当てはめたワイブル、'
            '対数正規、ガンマ分布のオーバーレイを示す。')

        ms_results_body = add_paragraph_after(doc, results_insert_ref, None)
        replace_paragraph_text_fully(ms_results_body, ms_results_text)

        add_heading_after(doc, results_insert_ref, 'モデル選択の結果', level=2)
        print("Inserted: Model selection results + figure")

    # =========================================================================
    # INSERT: Papacy comparison (before elective sub-analysis)
    # =========================================================================
    sa1_ref = None
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt.startswith('サブ解析1：選挙制') and p.style.name.startswith('Heading'):
            sa1_ref = doc.paragraphs[i - 1]
            break

    if sa1_ref is not None:
        papacy_text = (
            f'査読者1の提案に従い、1000年以前の教皇庁（N = {papacy["pre1000"]["n"]}）の'
            f'補足分析を実施し、主解析に含まれる1000年以降（N = {papacy["post1000"]["n"]}）と比較した。'
            f'1000年以前の教皇庁はk = {papacy["pre1000"]["k"]:.4f}'
            f'（95% CI [{papacy["pre1000"]["ci"][0]:.4f}, {papacy["pre1000"]["ci"][1]:.4f}]）であり、'
            f'1に近く、近似的に指数的（無記憶的）な継承動態を示唆する。'
            f'これはコンクラーヴェ以前の教皇選出が在位期間に関して本質的にランダムであったという仮説と一致する。'
            f'1000年以降の教皇庁（k = {papacy["post1000"]["k"]:.4f}、'
            f'95% CI [{papacy["post1000"]["ci"][0]:.4f}, {papacy["post1000"]["ci"][1]:.4f}]）は'
            f'わずかに高い形状パラメータを示し、コンクラーヴェ選挙制度の安定化効果と整合的である。'
            f'二標本コルモゴロフ・スミルノフ検定（p = {papacy["ks_2samp_p"]:.4f}）および'
            f'マン・ホイットニーU検定（p = {papacy["mw_p"]:.4f}）は0.05水準で二期間間に'
            f'統計的に有意な差がないことを示す。図10に両期間の生存関数、ワイブル確率プロット、'
            f'在位期間ヒストグラムを示す。'
        )

        papacy_fig = add_paragraph_after(doc, sa1_ref, None)
        replace_paragraph_text_fully(papacy_fig,
            '図10. 教皇庁：1000年以前 vs 1000年以降の比較。(A) カプラン・マイヤー生存関数と'
            'ワイブル当てはめ曲線。(B) ワイブル確率プロット。(C) 在位期間ヒストグラムと'
            'ワイブルPDFの当てはめ。')
        papacy_fig.runs[0].italic = True

        fig_path = os.path.join(FIG_DIR, 'fig_papacy_comparison.png')
        if os.path.exists(fig_path):
            img_p = OxmlElement('w:p')
            papacy_fig._element.addprevious(img_p)
            from docx.text.paragraph import Paragraph
            img_para = Paragraph(img_p, doc)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(fig_path, width=Inches(6.0))

        papacy_body = add_paragraph_after(doc, sa1_ref, None)
        replace_paragraph_text_fully(papacy_body, papacy_text)

        add_heading_after(doc, sa1_ref, '教皇庁：1000年以前 vs 1000年以降の比較', level=2)
        print("Inserted: Papacy comparison + figure")

    # =========================================================================
    # INSERT: Sample size assessment (before Discussion)
    # =========================================================================
    discussion_ref = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '考察' and p.style.name == 'Heading 1':
            discussion_ref = doc.paragraphs[i - 1]
            break

    if discussion_ref is not None:
        crossing_list = '、'.join(ss['crossing'])
        ss_text = (
            f'図12にワイブル形状パラメータkのサンプルサイズと信頼区間幅の関係を示す。'
            f'予想通り、信頼区間幅は概ね1/sqrt(N)に比例して減少し、'
            f'最小サンプルの4政体（N < 15）は信頼区間幅が1.0を超える。'
            f'{n_total}政体中{ss["n_crossing"]}政体の95%信頼区間がk = 1を跨ぐ：'
            f'{crossing_list}。'
            f'これらの政体については、定性的な体制分類（クーデター多発型 vs 安定型）は慎重に'
            f'解釈すべきであり、データは無記憶（指数的）過程との区別に不十分である。'
        )

        ss_fig = add_paragraph_after(doc, discussion_ref, None)
        replace_paragraph_text_fully(ss_fig,
            '図12. サンプルサイズ評価。(A) 全31政体のN vs 95%信頼区間幅'
            '（赤＝CIがk=1を跨ぐ、青＝跨がない）。N = 15の垂直破線は定性的解釈の最小閾値。'
            '(B) 政体別95%信頼区間のフォレストプロット。')
        ss_fig.runs[0].italic = True

        fig_path = os.path.join(FIG_DIR, 'fig_sample_size_analysis.png')
        if os.path.exists(fig_path):
            img_p = OxmlElement('w:p')
            ss_fig._element.addprevious(img_p)
            from docx.text.paragraph import Paragraph
            img_para = Paragraph(img_p, doc)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(fig_path, width=Inches(6.0))

        ss_body = add_paragraph_after(doc, discussion_ref, None)
        replace_paragraph_text_fully(ss_body, ss_text)

        add_heading_after(doc, discussion_ref, 'サンプルサイズ評価', level=2)
        print("Inserted: Sample size assessment + figure")

    # =========================================================================
    # INSERT: Saleh comparison in Discussion
    # =========================================================================
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith('継承メカニズム別のサブ解析'):
            saleh_text = (
                f'ローマ帝国に関する本研究の結果は、Saleh（2019）の先駆的研究との比較に値する。'
                f'Salehはローマ帝国の在位期間に2つのワイブル分布の混合を当てはめ、'
                f'初期の暗殺リスクと後期の摩耗を反映するバスタブ型ハザード関数を得た。'
                f'本研究の単一ワイブル当てはめ（k = {saleh["single_k"]:.3f}）は'
                f'ローマ帝国の全般的なクーデター多発的性格を捉えるが、二峰性ハザード構造を解像しない。'
                f'これに対処するため、ローマデータに独自の混合ワイブルを当てはめた：'
                f'混合モデル（k1 = {saleh["mix_k1"]:.3f}、k2 = {saleh["mix_k2"]:.3f}、'
                f'混合割合π = {saleh["mix_pi"]:.3f}）はAIC = {saleh["mix_aic"]:.1f}を達成し、'
                f'単一モデルのAIC = {saleh["single_aic"]:.1f}を上回る一方、'
                f'BICは単一モデルを支持する（{saleh["single_bic"]:.1f} vs {saleh["mix_bic"]:.1f}）。'
                f'主要な方法論的差異は、Salehが暴力的死亡までの時間を特異的に分析したのに対し、'
                f'本研究は終了原因にかかわらず総在位期間を分析している点にある。'
                f'文化横断的比較（本研究の主目的）には、単一の形状パラメータkが簡潔で普遍的に'
                f'適用可能な要約を提供する。図11に単一・混合当てはめの比較を示す。'
            )

            saleh_fig = add_paragraph_after(doc, p, None)
            replace_paragraph_text_fully(saleh_fig,
                '図11. ローマ帝国：単一 vs 混合ワイブル比較。(A) 確率密度関数。'
                '(B) ハザード関数：単一モデルは単調減少ハザード（k < 1）、'
                '混合モデルはSaleh（2019）の知見と一致するバスタブ曲線を示す。')
            saleh_fig.runs[0].italic = True

            fig_path = os.path.join(FIG_DIR, 'fig_roman_saleh_comparison.png')
            if os.path.exists(fig_path):
                img_p = OxmlElement('w:p')
                saleh_fig._element.addprevious(img_p)
                from docx.text.paragraph import Paragraph
                img_para = Paragraph(img_p, doc)
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(fig_path, width=Inches(6.0))

            saleh_body = add_paragraph_after(doc, p, None)
            replace_paragraph_text_fully(saleh_body, saleh_text)

            add_heading_after(doc, p, 'Salehの混合ワイブルモデルとの比較', level=2)
            print("Inserted: Saleh comparison + figure")
            break

    # =========================================================================
    # INSERT: Parliamentary democracies limitation
    # =========================================================================
    if 'future_work' in para_map:
        ref = doc.paragraphs[para_map['future_work'] - 1]
        parl_text = (
            '第六に、データセットには現代の議会制民主主義（例：西欧の首相）が含まれていない。'
            '議会制における首相の在任期間は根本的に異なる終了メカニズム（不信任投票、'
            '党首選、固定選挙周期）を伴い、別個の分析的扱いを要する。'
            'この不在は現代の統治システムの網羅性に非対称性を生じさせており、将来の研究で対処すべきである。'
        )
        parl_body = add_paragraph_after(doc, ref, None)
        replace_paragraph_text_fully(parl_body, parl_text)
        print("Inserted: Parliamentary democracies limitation")

    # =========================================================================
    # POST-PROCESSING
    # =========================================================================
    for p in doc.paragraphs:
        txt = p.text
        if 'Palgrave Communications誌に掲載された研究で' in txt:
            new_txt = txt.replace('Palgrave Communications誌に掲載された研究で', '')
            replace_paragraph_text_fully(p, new_txt)
            print("POST: Fixed Palgrave inline mention")

        if '歴史的検証' in txt and '内部整合性' not in txt and '独立した外部検証' not in txt:
            if p.style.name.startswith('Heading') or '歴史的検証分析' in txt or '歴史的検証：' in txt:
                new_txt = txt.replace('歴史的検証', '内部整合性の検証')
                replace_paragraph_text_fully(p, new_txt)
                print(f"POST: Fixed '歴史的検証' in paragraph")

    # =========================================================================
    # Save
    # =========================================================================
    out_path = os.path.join(OUT_DIR, 'hassc_manuscript_ja_v4.docx')
    doc.save(out_path)
    print(f"\nSaved JA v4 to: {out_path}")
    print(f"JA V4 has {len(doc.paragraphs)} paragraphs (v3 had 130)")
    return out_path


if __name__ == '__main__':
    generate_ja_v4()
    print("\nDone!")
