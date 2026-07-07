#!/usr/bin/env python3
"""
Post-process JA v4 manuscript:
1. Renumber figures/tables in order of first appearance
2. Convert author-date citations to Vancouver numbered superscript
3. Reorder reference list by citation order
"""

import os
import re
import copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.join(SCRIPT_DIR, '..')
OUT_DIR = os.path.join(ROOT_DIR, 'manuscript')

# JA figure appearance order: 1, 10, 2, 3, 4, 5, 9, S1, 6, 7, 8, 12, 11
JA_FIG_RENAME = {
    '1': '1',
    '10': '2',
    '2': '3',
    '3': '4',
    '4': '5',
    '5': '6',
    '9': '7',
    '6': '8',
    '7': '9',
    '8': '10',
    '12': '11',
    '11': '12',
}

# JA tables are already in order: 1, 2, 3, 4, 5 → no renaming needed

# Citation order by first appearance in JA text
JA_CITATION_ORDER = [
    # P9 Abstract
    ('Saleh', 'Saleh'),
    # P13
    ('Lawless, 2003', 'Lawless'),
    ('Weibull, 1951', 'Weibull, W. (1951)'),
    ('Abernethy, 2004', 'Abernethy'),
    ("O'Connor and Kleyner, 2012", "O'Connor"),
    # P15
    ('美濃部, 1912', '美濃部達吉 (1912)'),
    ('美濃部, 1923', '美濃部達吉 (1923)'),
    ('Miller, 1965', 'Miller'),
    ('Banno, 2001', 'Banno'),
    # P16
    ('Taagepera, 1979', 'Taagepera'),
    ('Turchin, 2003', 'Turchin, P. (2003)'),
    ('Turchin and Nefedov, 2009', 'Turchin, P. and Nefedov'),
    ('Clauset, 2018', 'Clauset'),
    ('Murdock, 1967', 'Murdock'),
    ('Olsson and Paik, 2016', 'Olsson'),
    # P29
    ('Burnham and Anderson, 2002', 'Burnham'),
    # P33
    ('Kaplan and Meier, 1958', 'Kaplan'),
    ('Cox, 1972', 'Cox'),
    # P72
    ('Goemans et al., 2009', 'Goemans'),
    # P105/106 (Discussion)
    ('Davies, 2005', 'Davies'),
    ('Svolik, 2012', 'Svolik'),
    # P115
    ('Scheidel, 2017', 'Scheidel'),
    ('Pinker, 2011', 'Pinker'),
    ('Goldstein, 2011', 'Goldstein'),
    ('North et al., 2009', 'North'),
    ('Kennedy, 1987', 'Kennedy'),
]

JA_CITE_NUM = {}
for idx, (key, _) in enumerate(JA_CITATION_ORDER):
    JA_CITE_NUM[key] = idx + 1


def replace_paragraph_text_fully(paragraph, new_text):
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        run = paragraph.add_run(new_text)
        run.font.size = Pt(11)


def rebuild_paragraph_with_superscript(paragraph, text):
    first_run_font = None
    if paragraph.runs:
        fr = paragraph.runs[0]
        first_run_font = {
            'size': fr.font.size,
            'name': fr.font.name,
        }

    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)
    if paragraph.runs:
        paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)

    parts = re.split(r'(\{[^}]+\})', text)

    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if first_run_font:
            if first_run_font['size']:
                run.font.size = first_run_font['size']
            if first_run_font['name']:
                run.font.name = first_run_font['name']

        if part.startswith('{') and part.endswith('}'):
            run.text = part[1:-1]
            run.font.superscript = True
        else:
            run.text = part


def rename_ja_figures(doc):
    for p in doc.paragraphs:
        txt = p.text
        changed = False
        new_txt = txt

        for old_num, new_num in JA_FIG_RENAME.items():
            if old_num == new_num:
                continue
            # 図N patterns (various suffixes)
            for suffix in ['.', ' ', '。', 'に', 'は', 'を', 'の', 'が', '、', '）', ')']:
                old_pat = f'図{old_num}{suffix}'
                new_pat = f'__FIG_{new_num}__{suffix}'
                if old_pat in new_txt:
                    new_txt = new_txt.replace(old_pat, new_pat)
                    changed = True
            # Figure N patterns (English in JA doc captions)
            for suffix in ['.', ' ', ',', ')', ';']:
                old_pat = f'Figure {old_num}{suffix}'
                new_pat = f'__FIG_{new_num}__{suffix}'
                if old_pat in new_txt:
                    new_txt = new_txt.replace(old_pat, new_pat)
                    changed = True

        if changed:
            replace_paragraph_text_fully(p, new_txt)

    # Pass 2: Replace placeholders
    for p in doc.paragraphs:
        txt = p.text
        if '__FIG_' in txt:
            new_txt = re.sub(r'__FIG_(\d+)__', r'図\1', txt)
            replace_paragraph_text_fully(p, new_txt)


def convert_ja_citations(doc):
    """Convert author-date citations to Vancouver numbered superscript in JA text."""
    # Define all replacement patterns (longest first)
    replacements = [
        # Multi-citation parenthetical (Japanese brackets)
        ('（Weibull, 1951; Abernethy, 2004）', '{3,4}'),
        ('（Miller, 1965; Banno, 2001）', '{8,9}'),
        ('（Turchin, 2003; Turchin and Nefedov, 2009）', '{11,12}'),
        ('（Murdock, 1967; Olsson and Paik, 2016）', '{14,15}'),
        ('（Pinker, 2011; Goldstein, 2011）', '{23,24}'),
        # Single parenthetical (Japanese brackets)
        ('（Lawless, 2003）', '{2}'),
        ("（O'Connor and Kleyner, 2012）", '{5}'),
        ('（美濃部, 1912, 1923）', '{6,7}'),
        ('（Taagepera, 1979）', '{10}'),
        ('（Clauset, 2018）', '{13}'),
        ('（Burnham and Anderson, 2002）', '{16}'),
        ('（Kaplan and Meier, 1958）', '{17}'),
        ('（Goemans et al., 2009）', '{19}'),
        ('（Davies, 2005）', '{20}'),
        ('（Svolik, 2012）', '{21}'),
        ('（Scheidel, 2017）', '{22}'),
        ('（North et al., 2009）', '{25}'),
        ('（Kennedy, 1987）', '{26}'),
        # Narrative (Japanese)
        ('Saleh（2019）', 'Saleh{1}'),
        ('Cox（1972）', 'Cox{18}'),
        ('Weibull（1951）', 'Weibull{3}'),
    ]

    for p in doc.paragraphs:
        if p.style.name == 'List Bullet':
            continue
        txt = p.text
        if not txt.strip():
            continue

        new_txt = txt
        for old, new in sorted(replacements, key=lambda x: -len(x[0])):
            new_txt = new_txt.replace(old, new)

        if new_txt != txt:
            rebuild_paragraph_with_superscript(p, new_txt)


def reorder_ja_references(doc):
    ref_paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'List Bullet' and p.text.strip():
            ref_paragraphs.append((i, p.text.strip()))

    if not ref_paragraphs:
        return

    ordered_refs = []
    used_indices = set()

    for _, ref_prefix in JA_CITATION_ORDER:
        for idx, (para_idx, txt) in enumerate(ref_paragraphs):
            if idx in used_indices:
                continue
            if txt.startswith(ref_prefix):
                ordered_refs.append((para_idx, txt))
                used_indices.add(idx)
                break

    for idx, (para_idx, txt) in enumerate(ref_paragraphs):
        if idx not in used_indices:
            print(f"WARNING: Orphan JA reference: {txt[:60]}")
            ordered_refs.append((para_idx, txt))

    ref_para_objects = [doc.paragraphs[pi] for pi, _ in ref_paragraphs]
    for i, (ref_obj, (_, new_text)) in enumerate(zip(ref_para_objects, ordered_refs)):
        numbered_text = f'{i + 1}. {new_text}'
        replace_paragraph_text_fully(ref_obj, numbered_text)


def verify_ja(doc):
    fig_order = []
    for p in doc.paragraphs:
        figs = re.findall(r'図(\d+)', p.text)
        for f in figs:
            if f not in fig_order:
                fig_order.append(f)
    print(f"JA Figures: {fig_order}")
    expected = [str(i) for i in range(1, 13)]
    if fig_order == expected:
        print("  ✓ JA figures correctly numbered 1-12")
    else:
        print(f"  ✗ Expected {expected}, got {fig_order}")


def add_missing_references_ja(doc):
    """Add references that are cited in text but missing from the reference list."""
    # Check for Burnham
    has_burnham = any(
        p.style.name == 'List Bullet' and 'Burnham' in p.text
        for p in doc.paragraphs
    )
    if not has_burnham:
        # Find last List Bullet paragraph and add after it
        last_ref_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.style.name == 'List Bullet' and p.text.strip():
                last_ref_idx = i
        if last_ref_idx is not None:
            from docx.oxml import OxmlElement
            ref_text = 'Burnham, K.P. and Anderson, D.R. (2002). Model Selection and Multimodel Inference: A Practical Information-Theoretic Approach (2nd ed.). Springer.'
            new_p = OxmlElement('w:p')
            doc.paragraphs[last_ref_idx]._element.addnext(new_p)
            from docx.text.paragraph import Paragraph
            new_para = Paragraph(new_p, doc)
            new_para.style = doc.paragraphs[last_ref_idx].style
            run = new_para.add_run(ref_text)
            run.font.size = Pt(11)
            print("Added missing Burnham and Anderson (2002) to JA reference list")


def postprocess_ja(input_path, output_path):
    print(f"Reading: {input_path}")
    doc = Document(input_path)
    print(f"Document has {len(doc.paragraphs)} paragraphs")

    add_missing_references_ja(doc)
    rename_ja_figures(doc)
    convert_ja_citations(doc)
    reorder_ja_references(doc)

    doc.save(output_path)
    print(f"Saved to: {output_path}")

    verify_doc = Document(output_path)
    verify_ja(verify_doc)


if __name__ == '__main__':
    ja_v4 = os.path.join(OUT_DIR, 'hassc_manuscript_ja_v4.docx')
    ja_v4_out = os.path.join(OUT_DIR, 'hassc_manuscript_ja_v4_final.docx')
    postprocess_ja(ja_v4, ja_v4_out)
