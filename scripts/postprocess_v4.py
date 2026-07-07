#!/usr/bin/env python3
"""
Post-process v4 manuscripts:
1. Renumber figures/tables in order of first appearance
2. Convert author-date citations to Vancouver numbered superscript
3. Reorder reference list by citation order
4. Fix structural issues (parliamentary democracies paragraph placement)
5. Polish English in new sections
"""

import os
import re
import copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.join(SCRIPT_DIR, '..')
OUT_DIR = os.path.join(ROOT_DIR, 'manuscript')

# ============================================================================
# FIGURE / TABLE RENUMBERING
# ============================================================================

# Map old figure numbers → new (in order of appearance from v4 text analysis)
# Appearance order: 1, 9, S1, 6, 7, 8, 10, 2, 3, 4, 5, 12, 11
FIG_RENAME = {
    '1': '1',    # unchanged
    '9': '2',    # model selection
    '6': '3',    # internal consistency
    '7': '4',    # accession age
    '8': '5',    # sensitivity
    '10': '6',   # papacy
    '2': '7',    # elective map
    '3': '8',    # elective comparison
    '4': '9',    # era stratification
    '5': '10',   # box plots
    '12': '11',  # sample size
    '11': '12',  # Saleh comparison
}

# Map old table numbers → new (in order of appearance)
# Appearance order: A1, 1, 3, 4, 5, 2
TBL_RENAME = {
    '1': '1',    # unchanged
    '3': '2',    # violent succession
    '4': '3',    # accession age
    '5': '4',    # sensitivity
    '2': '5',    # era stratification
}

# ============================================================================
# CITATION MAPPING: author-date → Vancouver numbered
# ============================================================================

# Citations in order of first appearance in text (manually traced from v4)
# Each entry: (pattern_for_matching, reference_list_prefix)
CITATION_ORDER = [
    # P14 Introduction
    ('Lawless, 2003', 'Lawless'),
    ('Weibull, 1951', 'Weibull, W. (1951)'),
    ('Abernethy, 2004', 'Abernethy, R.B. (2004)'),
    ("O'Connor and Kleyner, 2012", "O'Connor, P.D.T."),
    # P15
    ('Saleh, 2019', 'Saleh'),  # also Saleh's (2019), Saleh (2019)
    # P16
    ('Minobe, 1912', 'Minobe, T. (1912)'),
    ('Minobe, 1923', 'Minobe, T. (1923)'),
    ('Miller, 1965', 'Miller'),
    ('Banno, 2001', 'Banno'),
    # P17
    ('Taagepera, 1979', 'Taagepera'),
    ('Turchin, 2003', 'Turchin, P. (2003)'),
    ('Turchin and Nefedov, 2009', 'Turchin, P. and Nefedov'),
    ('Clauset, 2018', 'Clauset'),
    ('Murdock, 1967', 'Murdock'),
    ('Olsson and Paik, 2016', 'Olsson'),
    # P36 Model selection
    ('Burnham and Anderson, 2002', 'Burnham'),
    # P40 Statistical comparisons
    ('Kaplan and Meier, 1958', 'Kaplan'),
    ('Cox, 1972', 'Cox'),
    # P42 Internal consistency
    ('Goemans et al., 2009', 'Goemans'),
    # P122 Discussion accession age
    ('Davies, 2005', 'Davies'),
    # P124 Discussion elective
    ('Svolik, 2012', 'Svolik'),
    # P129 Era discussion
    ('Scheidel, 2017', 'Scheidel'),
    ('Pinker, 2011', 'Pinker'),
    ('Goldstein, 2011', 'Goldstein'),
    ('North et al., 2009', 'North'),
    ('Kennedy, 1987', 'Kennedy'),
]

# Build citation key → number mapping
CITE_NUM = {}
for idx, (key, _) in enumerate(CITATION_ORDER):
    CITE_NUM[key] = idx + 1

# ============================================================================
# ENGLISH POLISH: replacement map for AI-sounding phrases in new sections
# ============================================================================

ENGLISH_POLISH = [
    # Model selection methods (P36)
    (
        'To assess whether the Weibull distribution provides the best available fit rather than merely an acceptable one, we conducted systematic model selection comparing four candidate distributions for positive, right-skewed survival data: Weibull, log-normal, gamma, and exponential.',
        'Whether the Weibull distribution provides the best available fit - not merely an acceptable one - was assessed through model selection across four candidate distributions for positive, right-skewed survival data: Weibull, log-normal, gamma, and exponential.'
    ),
    (
        'All distributions were fitted via MLE with location parameter fixed at zero. Model comparison was performed using the Akaike Information Criterion (AIC) and the Bayesian Information Criterion (BIC).',
        'All distributions were fitted by maximum likelihood with location fixed at zero. Models were compared using the Akaike Information Criterion (AIC) and the Bayesian Information Criterion (BIC).'
    ),
    # Sample size methods (P38)
    (
        'Polities with sample sizes below N = 15 (Aztec Empire N = 9, Thailand Chakri N = 10, China Qing N = 11, Inca Empire N = 12) yield wide confidence intervals that preclude reliable qualitative interpretation.',
        'Four polities fall below a practical threshold of N = 15: the Aztec Empire (N = 9), Thailand Chakri (N = 10), China Qing (N = 11), and the Inca Empire (N = 12). Their confidence intervals are too wide for reliable qualitative interpretation.'
    ),
    (
        'These polities are included in Table 1 for completeness but are flagged and excluded from qualitative regime classification. Additionally, polities whose 95% CI for k spans the k = 1 threshold are noted as having indeterminate stability regime classification.',
        'These cases remain in Table 1 for completeness but are flagged and excluded from qualitative regime classification. Polities whose 95% CI for k straddles the k = 1 boundary are likewise noted as indeterminate.'
    ),
    # Model selection results (P53) - break up long paragraph
    (
        'Systematic model selection comparing Weibull, log-normal, gamma, and exponential distributions was conducted for all 31 polities.',
        'We compared the Weibull, log-normal, gamma, and exponential distributions across all 31 polities.'
    ),
    (
        'Critically, the Weibull fell within ΔAIC < 2 of the best model for 26 of 31 polities, indicating that it provides a statistically competitive fit for the large majority of the dataset.',
        'The Weibull fell within ΔAIC < 2 of the best model for 26 of 31 polities, confirming a competitive fit across the large majority of the dataset.'
    ),
    (
        'These results confirm that the Weibull distribution provides a defensible analytical framework for cross-polity comparison, while acknowledging that it is not universally the best-fitting model.',
        'The Weibull thus provides a defensible framework for cross-polity comparison, though it is not universally the best-fitting model.'
    ),
    (
        'The interpretive advantage of the shape parameter k - its direct mapping onto hazard dynamics via the bathtub curve framework - justifies its use as the primary metric even where alternative distributions achieve marginally better statistical fit.',
        'Because k maps directly onto hazard dynamics through the bathtub-curve framework, it remains the preferred metric even for polities where an alternative distribution achieves a marginally better statistical fit.'
    ),
    # S1 reference (P56)
    (
        'Supplementary Figure S1 presents empirical Kaplan-Meier survival curves overlaid with fitted Weibull, log-normal, and gamma distributions for all 31 polities, allowing visual assessment of fit quality. As Reviewer 1 anticipated, some polities (notably those with small N or heavy-tailed durations) show visible deviations at the extremes; these are discussed below as informative departures that enrich the interpretation rather than constituting weaknesses.',
        'Supplementary Figure S1 overlays empirical Kaplan-Meier survival curves with the fitted Weibull, log-normal, and gamma distributions for each of the 31 polities. Some polities - particularly those with small N or heavy-tailed durations - show departures from the Weibull at the extremes; we treat these as informative rather than as weaknesses.'
    ),
    # Papacy comparison (P90)
    (
        'Following Reviewer 1\'s suggestion, we conducted a supplementary analysis of the pre-1000 CE papacy',
        'We also examined the pre-1000 CE papacy'
    ),
    (
        'This is consistent with the hypothesis that pre-conclave papal selection was essentially random with respect to tenure duration, dominated by aristocratic faction politics and imperial intervention rather than institutional mechanisms favouring stability.',
        'Pre-conclave papal selection appears to have been essentially random with respect to tenure, governed by aristocratic faction politics and imperial intervention rather than by institutions favouring stability.'
    ),
    (
        'The two-sample Kolmogorov-Smirnov test (p = 0.5629) and Mann-Whitney U test (p = 0.1347) indicate no statistically significant difference between the two periods at the 0.05 level, although the shift from near-exponential to mildly increasing hazard is directionally consistent with institutional stabilisation.',
        'Neither the two-sample Kolmogorov-Smirnov test (p = 0.5629) nor the Mann-Whitney U test (p = 0.1347) detects a significant difference at the 0.05 level, though the shift from near-exponential toward mildly increasing hazard is directionally consistent with institutional stabilisation.'
    ),
    # Saleh comparison (P126)
    (
        'Our results for the Roman Empire merit comparison with Saleh\'s (2019) pioneering study.',
        'The Roman Empire results invite direct comparison with Saleh\'s (2019) original study.'
    ),
    (
        'The key methodological difference is that Saleh analysed time-to-violent-death specifically, whereas we analyse total reign duration regardless of termination cause.',
        'A key methodological difference is that Saleh analysed time to violent death specifically, whereas we analyse total reign duration irrespective of how the reign ended.'
    ),
    (
        'For cross-polity comparison (our primary objective), the single shape parameter k provides a parsimonious and universally applicable summary, while mixture models may reveal finer within-polity dynamics where sample sizes permit.',
        'For cross-polity comparison - our primary objective - the single shape parameter k offers a parsimonious summary; mixture models can reveal finer within-polity structure where sample sizes permit.'
    ),
    # Sample size assessment (P115) - was in wrong section but also needs polish
    (
        'Figure 12 illustrates the relationship between sample size and confidence interval width for the Weibull shape parameter k. As expected, CI width decreases approximately as 1/sqrt(N), with the four smallest-sample polities (N < 15) exhibiting CI widths exceeding 1.0.',
        'Figure 12 plots confidence interval width against sample size for k. As expected, CI width falls roughly as 1/√N; the four polities with N < 15 all have CI widths above 1.0.'
    ),
    (
        'For these polities, qualitative regime classification (coup-prone vs. stable) should be treated with caution, as the data are insufficient to distinguish the polity\'s hazard dynamics from those of a memoryless (exponential) process.',
        'For these polities, the data cannot reliably distinguish the hazard pattern from an exponential (memoryless) process, and qualitative regime labels should be read with caution.'
    ),
    # Discussion general polish
    (
        'The conceptual framework underlying this analysis draws on two intellectual traditions.',
        'Two intellectual traditions underpin this analysis.'
    ),
    (
        'Our main finding - that the Weibull shape parameter varies substantially across polities (range: 0.851-2.721) and tracks intuitive notions of political stability - supports the use of k as a quantitative \'peacefulness index\'.',
        'The central finding - that k varies widely across polities (0.851-2.721) and tracks intuitive notions of political stability - supports its use as a quantitative \'peacefulness index\'.'
    ),
    # Discussion consistency section
    (
        'The internal consistency analysis demonstrates that k carries substantive historical meaning.',
        'The internal consistency analysis confirms that k is not merely a statistical abstraction.'
    ),
    # Era discussion
    (
        'The era-based sub-analysis reveals a striking temporal pattern: a monotonic increase in k from antiquity through the modern period, followed by a reversal in the contemporary era.',
        'The era-stratified results trace a clear temporal arc: k rises monotonically from antiquity through the modern period and then reverses in the contemporary era.'
    ),
    (
        'This trajectory is consistent with broader arguments about the decline of political violence over historical time (Pinker, 2011; Goldstein, 2011) and the role of institutional frameworks in constraining violence (North et al., 2009).',
        'This arc echoes broader arguments about the long-run decline of political violence (Pinker, 2011; Goldstein, 2011) and the constraining role of institutions (North et al., 2009).'
    ),
    # Limitations
    (
        'Several limitations warrant discussion.',
        'Several limitations should be acknowledged.'
    ),
    # Parliamentary democracies paragraph polish
    (
        'Sixth, the dataset does not include contemporary parliamentary democracies (e.g., Western European prime ministers). Prime ministerial tenures in parliamentary systems involve fundamentally different termination mechanisms - votes of confidence, party leadership contests, and fixed electoral cycles - that merit separate analytical treatment. Their absence creates an asymmetry in the coverage of modern governance systems that future work should address.',
        'Sixth, the dataset excludes contemporary parliamentary democracies (e.g., Western European prime ministers). Prime ministerial tenures end through votes of confidence, party leadership contests, and fixed electoral cycles - mechanisms fundamentally different from those governing monarchical or presidential succession. This omission introduces an asymmetry in coverage of modern governance that future work should address.'
    ),
]


def rename_figures_tables(doc):
    """Renumber figures and tables in order of first appearance using two-pass placeholder approach."""
    # Pass 1: Replace old numbers with placeholders
    for p in doc.paragraphs:
        txt = p.text
        changed = False
        new_txt = txt

        # Figure renaming (use __FIG_XX__ placeholders)
        for old_num, new_num in FIG_RENAME.items():
            if old_num == new_num:
                continue
            patterns = [
                (f'Figure {old_num}.', f'__FIG_{new_num}__.'),
                (f'Figure {old_num} ', f'__FIG_{new_num}__ '),
                (f'Figure {old_num},', f'__FIG_{new_num}__,'),
                (f'Figure {old_num})', f'__FIG_{new_num}__)'),
                (f'Figure {old_num};', f'__FIG_{new_num}__;'),
                (f'Figure {old_num}\n', f'__FIG_{new_num}__\n'),
                (f'図{old_num}.', f'__FIG_{new_num}__.'),
                (f'図{old_num} ', f'__FIG_{new_num}__ '),
                (f'図{old_num}。', f'__FIG_{new_num}__。'),
                (f'図{old_num}に', f'__FIG_{new_num}__に'),
            ]
            for old_pat, new_pat in patterns:
                if old_pat in new_txt:
                    new_txt = new_txt.replace(old_pat, new_pat)
                    changed = True

        # Table renaming (use __TBL_XX__ placeholders)
        for old_num, new_num in TBL_RENAME.items():
            if old_num == new_num:
                continue
            patterns = [
                (f'Table {old_num}.', f'__TBL_{new_num}__.'),
                (f'Table {old_num} ', f'__TBL_{new_num}__ '),
                (f'Table {old_num},', f'__TBL_{new_num}__,'),
                (f'Table {old_num})', f'__TBL_{new_num}__)'),
                (f'Table {old_num};', f'__TBL_{new_num}__;'),
                (f'Table {old_num}\n', f'__TBL_{new_num}__\n'),
                (f'表{old_num}.', f'__TBL_{new_num}__.'),
                (f'表{old_num} ', f'__TBL_{new_num}__ '),
                (f'表{old_num}に', f'__TBL_{new_num}__に'),
                (f'表{old_num}の', f'__TBL_{new_num}__の'),
            ]
            for old_pat, new_pat in patterns:
                if old_pat in new_txt:
                    new_txt = new_txt.replace(old_pat, new_pat)
                    changed = True

        if changed:
            replace_paragraph_text_fully(p, new_txt)

    # Pass 2: Replace placeholders with final numbers
    for p in doc.paragraphs:
        txt = p.text
        if '__FIG_' in txt or '__TBL_' in txt:
            new_txt = re.sub(r'__FIG_(\d+)__', r'Figure \1', txt)
            new_txt = re.sub(r'__TBL_(\d+)__', r'Table \1', new_txt)
            replace_paragraph_text_fully(p, new_txt)


def convert_citations_vancouver(doc):
    """Convert author-date citations to Vancouver numbered superscript style."""
    # Build replacement patterns ordered from longest to shortest to avoid partial matches
    replacements = []

    # Handle various citation patterns for each reference
    for cite_key, num in CITE_NUM.items():
        author = cite_key.split(',')[0].strip()
        year = cite_key.split(',')[-1].strip()

        # Parenthetical: (Author, Year) or (Author and Author, Year)
        replacements.append((f'({cite_key})', f'{{{num}}}'))

        # Possessive: Author's (Year)
        replacements.append((f"{author}'s ({year})", f"{author}'s{{{num}}}"))

        # Narrative: Author (Year)
        replacements.append((f'{author} ({year})', f'{author}{{{num}}}'))

    # Handle multi-citation patterns
    # (Weibull, 1951; Abernethy, 2004) → {2,3}
    # (Minobe, 1912, 1923) → {6,7}
    # (Turchin, 2003; Turchin and Nefedov, 2009) → {11,12}
    # etc.
    multi_cites = [
        ('(Lawless, 2003)', '{1}'),
        ('(Weibull, 1951; Abernethy, 2004)', '{2,3}'),
        ('(Weibull, 1951)', '{2}'),
        ("(O'Connor and Kleyner, 2012)", '{4}'),
        ('(Minobe, 1912, 1923)', '{6,7}'),
        ('(Minobe, 1912)', '{6}'),
        ('(Minobe, 1923)', '{7}'),
        ('(Miller, 1965; Banno, 2001)', '{8,9}'),
        ('(Miller, 1965)', '{8}'),
        ('(Banno, 2001)', '{9}'),
        ('(Taagepera, 1979)', '{10}'),
        ('(Turchin, 2003; Turchin and Nefedov, 2009)', '{11,12}'),
        ('(Turchin, 2003)', '{11}'),
        ('(Turchin and Nefedov, 2009)', '{12}'),
        ('(Clauset, 2018)', '{13}'),
        ('(Murdock, 1967; Olsson and Paik, 2016)', '{14,15}'),
        ('(Murdock, 1967)', '{14}'),
        ('(Olsson and Paik, 2016)', '{15}'),
        ('(Burnham and Anderson, 2002)', '{16}'),
        ('(Kaplan and Meier, 1958)', '{17}'),
        ('(Cox, 1972)', '{18}'),
        ('(Goemans et al., 2009)', '{19}'),
        ('(Davies, 2005)', '{20}'),
        ('(Svolik, 2012)', '{21}'),
        ('(Scheidel, 2017)', '{22}'),
        ('(Pinker, 2011; Goldstein, 2011)', '{23,24}'),
        ('(Pinker, 2011)', '{23}'),
        ('(Goldstein, 2011)', '{24}'),
        ('(North et al., 2009)', '{25}'),
        ('(Kennedy, 1987)', '{26}'),
    ]

    for p in doc.paragraphs:
        if p.style.name == 'List Bullet':
            continue  # Don't touch reference list
        txt = p.text
        if not txt.strip():
            continue

        new_txt = txt

        # Apply multi-citation replacements first (longer patterns first)
        for old, new in sorted(multi_cites, key=lambda x: -len(x[0])):
            new_txt = new_txt.replace(old, new)

        # Handle narrative citations: "Saleh's (2019)" → "Saleh's{5}"
        # "Saleh (2019)" → "Saleh{5}"
        # "Weibull's (1951)" → "Weibull's{2}"
        narrative_patterns = [
            ("Saleh's (2019)", "Saleh's{5}"),
            ("Saleh (2019)", "Saleh{5}"),
            ("Weibull's (1951)", "Weibull's{2}"),
            ("Weibull (1951)", "Weibull{2}"),
            ("Minobe, 1912, 1923", "Minobe{6,7}"),
            ("Lawless, 2003", "Lawless{1}"),
            ("Abernethy, 2004", "Abernethy{3}"),
        ]
        for old, new in narrative_patterns:
            new_txt = new_txt.replace(old, new)

        # Clean up remaining parenthetical citations that weren't caught
        # Pattern: (Author, Year) where Author starts with capital
        def replace_remaining_cite(match):
            full = match.group(0)
            inner = match.group(1)
            # Try to find this in our map
            for ck, cn in CITE_NUM.items():
                if ck in inner:
                    return '{' + str(cn) + '}'
            return full  # Leave unchanged if not found

        new_txt = re.sub(
            r'\(([A-Z][a-z]+(?:\s+(?:and|et\s+al\.?)\s+[A-Z][a-z]+)?,\s*\d{4}(?:;\s*[A-Z][a-z]+(?:\s+(?:and|et\s+al\.?)\s+[A-Z][a-z]+)?,\s*\d{4})*)\)',
            replace_remaining_cite,
            new_txt
        )

        if new_txt != txt:
            # Now we need to rebuild the paragraph with superscript runs for {N} markers
            rebuild_paragraph_with_superscript(p, new_txt)


def rebuild_paragraph_with_superscript(paragraph, text):
    """Split text on {N} markers and create runs with superscript for citation numbers."""
    # Preserve the formatting of the first run
    first_run_font = None
    if paragraph.runs:
        fr = paragraph.runs[0]
        first_run_font = {
            'size': fr.font.size,
            'bold': fr.font.bold,
            'italic': fr.font.italic,
            'name': fr.font.name,
        }

    # Clear existing runs
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)
    if paragraph.runs:
        paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)

    # Split text on {N} markers
    parts = re.split(r'(\{[^}]+\})', text)

    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if first_run_font:
            if first_run_font['size']:
                run.font.size = first_run_font['size']
            if first_run_font['name']:
                run.font.name = first_run_font['name']

        if part.startswith('{') and part.endswith('}'):
            # This is a citation number - make it superscript
            run.text = part[1:-1]  # Remove braces
            run.font.superscript = True
        else:
            run.text = part
            # Preserve italic for figure captions
            if first_run_font and first_run_font['italic']:
                run.font.italic = True


def reorder_references(doc):
    """Reorder reference list to match Vancouver citation order."""
    # Find reference list paragraphs
    ref_paragraphs = []
    ref_start_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'List Bullet' and p.text.strip():
            if ref_start_idx is None:
                ref_start_idx = i
            ref_paragraphs.append((i, p.text.strip()))

    if not ref_paragraphs:
        return

    # Match references to citation order
    ordered_refs = []
    used_indices = set()

    for _, ref_prefix in CITATION_ORDER:
        for idx, (para_idx, txt) in enumerate(ref_paragraphs):
            if idx in used_indices:
                continue
            if txt.startswith(ref_prefix):
                ordered_refs.append((para_idx, txt))
                used_indices.add(idx)
                break

    # Add any remaining references not matched (orphans - should flag)
    for idx, (para_idx, txt) in enumerate(ref_paragraphs):
        if idx not in used_indices:
            print(f"WARNING: Orphan reference not cited in text: {txt[:80]}")
            # Still include it at the end
            ordered_refs.append((para_idx, txt))

    # Now rewrite reference list with Vancouver numbering
    ref_para_objects = [doc.paragraphs[pi] for pi, _ in ref_paragraphs]

    for i, (ref_obj, (_, new_text)) in enumerate(zip(ref_para_objects, ordered_refs)):
        numbered_text = f'{i + 1}. {new_text}'
        replace_paragraph_text_fully(ref_obj, numbered_text)


def apply_english_polish(doc):
    """Apply targeted English improvements to remove AI-sounding language."""
    for p in doc.paragraphs:
        txt = p.text
        if not txt.strip():
            continue

        new_txt = txt
        for old, new in ENGLISH_POLISH:
            if old in new_txt:
                new_txt = new_txt.replace(old, new)

        if new_txt != txt:
            replace_paragraph_text_fully(p, new_txt)


def fix_structure(doc):
    """Fix structural issues:
    - Move parliamentary democracies paragraph from Sample size assessment to Limitations
    """
    # Find the parliamentary democracies paragraph and move it
    parl_idx = None
    limitations_last_idx = None

    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt.startswith('Sixth, the dataset') and 'parliamentary democracies' in txt:
            parl_idx = i
        elif txt.startswith('Sixth, the dataset excludes') and 'parliamentary democracies' in txt:
            parl_idx = i
        # Find the last paragraph before Conclusion that's in the Limitations section
        if txt.startswith('Future work could extend') or txt.startswith('Future work should extend'):
            limitations_last_idx = i

    if parl_idx is not None and limitations_last_idx is not None:
        # Check if it's currently in the wrong place (under Sample size assessment)
        # by checking if there's a "Sample size assessment" heading before it
        in_wrong_place = False
        for j in range(parl_idx - 1, max(0, parl_idx - 5), -1):
            if 'Sample size' in doc.paragraphs[j].text and doc.paragraphs[j].style.name.startswith('Heading'):
                in_wrong_place = True
                break

        if in_wrong_place and parl_idx != limitations_last_idx:
            # Remove from current location
            parl_text = doc.paragraphs[parl_idx].text
            doc.paragraphs[parl_idx]._element.getparent().remove(doc.paragraphs[parl_idx]._element)
            print(f"Moved parliamentary democracies paragraph to Limitations section")

            # It's now removed; we need to re-insert before the Future work paragraph
            # After removal, indices shifted, so re-find
            for i, p in enumerate(doc.paragraphs):
                if p.text.startswith('Future work could extend') or p.text.startswith('Future work should extend'):
                    # Insert before this paragraph
                    new_p = OxmlElement('w:p')
                    p._element.addprevious(new_p)
                    from docx.text.paragraph import Paragraph
                    new_para = Paragraph(new_p, doc)
                    run = new_para.add_run(parl_text)
                    run.font.size = Pt(11)
                    break


def replace_paragraph_text_fully(paragraph, new_text):
    """Replace all text in paragraph, keeping formatting of first run."""
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        run = paragraph.add_run(new_text)
        run.font.size = Pt(11)


def add_davies_citation(doc):
    """Add inline citation for Davies (2005) which is currently an orphan reference."""
    for p in doc.paragraphs:
        if 'rex regnat et non gubernat' in p.text and 'Davies' not in p.text:
            old = 'the Polish-Lithuanian Commonwealth\'s Golden Liberty'
            new = 'the Polish-Lithuanian Commonwealth\'s Golden Liberty (Davies, 2005)'
            txt = p.text.replace(old, new)
            replace_paragraph_text_fully(p, txt)
            print("Added Davies (2005) citation inline")
            break


def verify_completeness(doc):
    """Verify all figures, tables are cited and in order."""
    fig_order = []
    tbl_order = []

    for i, p in enumerate(doc.paragraphs):
        txt = p.text
        if not txt.strip():
            continue

        figs = re.findall(r'Figure\s+(\d+)', txt)
        for f in figs:
            if f not in fig_order:
                fig_order.append(f)

        tbls = re.findall(r'Table\s+(\d+)', txt)
        for t in tbls:
            if t not in tbl_order:
                tbl_order.append(t)

    print(f"\nFigures in order: {fig_order}")
    expected_figs = [str(i) for i in range(1, 13)]
    if fig_order == expected_figs:
        print("  ✓ Figures correctly numbered 1-12 in order")
    else:
        print(f"  ✗ Expected {expected_figs}, got {fig_order}")

    print(f"Tables in order: {tbl_order}")
    expected_tbls = ['1', '2', '3', '4', '5']
    # A1 may appear before or mixed
    main_tbls = [t for t in tbl_order if not t.startswith('A')]
    if main_tbls == expected_tbls:
        print("  ✓ Tables correctly numbered 1-5 in order")
    else:
        print(f"  ✗ Expected {expected_tbls}, got {main_tbls}")


def postprocess(input_path, output_path):
    print(f"Reading: {input_path}")
    doc = Document(input_path)
    print(f"Document has {len(doc.paragraphs)} paragraphs")

    # Step 0: Add Davies citation before converting to numbered style
    add_davies_citation(doc)

    # Step 1: Fix structure (move parliamentary democracies paragraph)
    fix_structure(doc)

    # Step 2: Apply English polish
    apply_english_polish(doc)

    # Step 3: Rename figures and tables
    rename_figures_tables(doc)

    # Step 4: Convert citations to Vancouver numbered
    convert_citations_vancouver(doc)

    # Step 5: Reorder reference list
    reorder_references(doc)

    # Save
    doc.save(output_path)
    print(f"\nSaved to: {output_path}")

    # Verify
    verify_doc = Document(output_path)
    verify_completeness(verify_doc)


if __name__ == '__main__':
    en_v4 = os.path.join(OUT_DIR, 'hassc_manuscript_en_v4.docx')
    en_v4_out = os.path.join(OUT_DIR, 'hassc_manuscript_en_v4_final.docx')
    postprocess(en_v4, en_v4_out)
