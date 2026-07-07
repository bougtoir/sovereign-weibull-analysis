#!/usr/bin/env python3
"""
Generate revised Japanese manuscript for HaSSC major revision.
"""

import os
import json
import re
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(SCRIPT_DIR, '..', 'figures')
OUT_DIR = os.path.join(SCRIPT_DIR, '..', 'manuscript')
os.makedirs(OUT_DIR, exist_ok=True)

with open(os.path.join(FIG_DIR, 'revision_results.json'), 'r') as f:
    RESULTS = json.load(f)


def add_citation_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(style=style)
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(9)
        else:
            run = p.add_run(part)
            run.font.size = Pt(11)
    return p


def generate_ja_manuscript():
    doc = Document()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('統治者在位期間のグローバルワイブル分析：\n'
                     '歴史的政治安定性の定量化に向けた信頼性工学的アプローチ')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[著者名]')
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[所属]')
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('責任著者: [氏名, メールアドレス]')
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Abstract
    doc.add_heading('要旨', level=1)
    add_citation_paragraph(doc,
        '信頼性工学における機械・電子部品の故障時間モデリングの基盤であるワイブル分布は、'
        'Salehによるローマ皇帝の在位期間分析を通じて政治史に応用された。{1} '
        '本研究では、このアプローチを5大陸31の主権政体に拡張し、約3,500年の記録された歴史と'
        '1,148件の個別在位期間を包括するグローバル比較分析を行った。'
        '最尤推定による2パラメータワイブル分布のフィッティングの結果、形状パラメータkは'
        '0.851（マムルーク朝）から2.721（米国大統領）の範囲を示し、クーデター多発体制'
        '（k < 1、減少ハザード）と安定体制（k > 1、増加ハザード）を識別する定量的「安定性指標」'
        'を提供することが明らかとなった。ワイブル、対数正規、ガンマ、指数分布のAIC・BICによる'
        '体系的モデル選択の結果、大多数の政体においてワイブル分布が競争力のあるフィット'
        '（ΔAIC < 2）を提供することが確認された。'
        'サブ分析では、選挙制と世襲制の間に有意差が認められ、古代から近代にかけてkが'
        '単調増加し、現代においてその傾向が反転する世俗的トレンドが明らかとなった。'
        'kと暴力的継承率の有意な負の相関により内部整合性が実証された。'
    )

    p = doc.add_paragraph()
    run = p.add_run('キーワード: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run('ワイブル分布; 信頼性工学; 政治的安定性; 在位期間; 比較史; '
                     '生存分析; モデル選択; 暴力的継承')
    run.font.size = Pt(11)

    # Introduction
    doc.add_heading('はじめに', level=1)
    add_citation_paragraph(doc,
        '信頼性工学手法の歴史的・社会的現象への応用は、定量的歴史学における新たなフロンティア'
        'を構成する。信頼性工学は主に機械・電子部品の寿命解析のために開発されたものであり、'
        '時間の関数としての「故障」確率をモデル化するための数学的に厳密な枠組みを提供する。{2,3} '
        '1951年にWaloddi Weibullにより導入されたワイブル分布は、{4} '
        'その形状パラメータkを通じて、減少（k < 1）、一定（k = 1）、増加（k > 1）の'
        'ハザード率をモデル化できる柔軟性から、この分野の標準ツールとなっている。'
    )
    add_citation_paragraph(doc,
        '機械的故障から政治的継承への概念的飛躍は、Salehにより先駆的に行われた。{1} '
        'Salehは69人のローマ皇帝の在位期間にワイブル分布をフィッティングし、2つのワイブル'
        '分布の混合モデルがローマ帝位継承の双峰性ハザード構造を捉えることを実証した。'
    )
    add_citation_paragraph(doc,
        '統治者在位期間を部品寿命と類似的に扱うことの理論的正当性は、美濃部達吉の'
        '「天皇機関説」に予想外の支持を見出す。{5,6} 美濃部は天皇を主権的個人としてではなく、'
        '国家機構の中の機能的構成要素として概念化した。{7}'
    )
    add_citation_paragraph(doc,
        '本研究の中心的仮説は、ワイブル形状パラメータkが堅牢で解釈可能かつ文化横断的に'
        '比較可能な政治的安定性の指標——ハザード関数の数学に基づいた「平和度指数」——を'
        '提供するということである。'
    )

    # Methods
    doc.add_heading('方法', level=1)

    doc.add_heading('データ収集と範囲', level=2)
    add_citation_paragraph(doc,
        '在位期間データは確立された歴史的資料から編纂し、各政体について複数の参考文献と'
        '照合して検証した。5つの地理的地域にわたる合計31の異なる政体、1,148件の個別在位期間が'
        '含まれる。データセット、変数定義、分析コードは'
        'https://github.com/bougtoir/sovereign-weibull-analysis で公開されている。'
    )

    doc.add_heading('継承メカニズムの分類', level=2)
    add_citation_paragraph(doc,
        '各政体は主要な継承メカニズムに基づき「選挙制」または「世襲制」に分類した。'
        '米国大統領は、憲法上の固定任期制が在位期間を機械的に制約するため、「任期制民主主義」'
        'として別カテゴリに分類した。'
    )

    doc.add_heading('モデル選択', level=2)
    add_citation_paragraph(doc,
        'ワイブル分布が最良のフィットを提供するかを評価するため、ワイブル、対数正規、'
        'ガンマ、指数の4つの候補分布について体系的モデル選択を実施した。モデル比較には'
        'AIC（赤池情報量基準）とBIC（ベイズ情報量基準）を用いた。{9}'
    )

    doc.add_heading('最小標本サイズの考慮', level=2)
    add_citation_paragraph(doc,
        '標本サイズがN = 15未満の政体（アステカ帝国N = 9、タイ・チャクリー朝N = 10、'
        '清朝N = 11、インカ帝国N = 12）は信頼区間が広く、定性的解釈が困難である。'
        'これらの政体は表1に参考値として含めるが、安定性レジーム分類からは除外する。'
    )

    doc.add_heading('内部整合性チェック', level=2)
    add_citation_paragraph(doc,
        'ワイブル形状パラメータが政治的安定性の経験的に意味のあるパターンを捉えているかを'
        '評価するため、暴力的継承率と内戦数を独立した歴史資料から編纂した。{11,12,13,14} '
        'ワイブル形状パラメータと暴力的継承率の両方が同一の歴史的記録に基づくため、'
        'この分析は独立した外部検証ではなく内部整合性の実証として位置付ける。'
    )

    # Results
    doc.add_heading('結果', level=1)

    doc.add_heading('主分析：31政体のグローバル比較', level=2)
    add_citation_paragraph(doc,
        '分析対象31政体の大多数でワイブル分布は適切なフィット（KS検定 p > 0.05）を示した。'
        '形状パラメータkは0.851（マムルーク朝）から2.721（米国大統領）の範囲であった。'
        '信頼区間がk = 1を跨ぐ政体については、レジーム分類を慎重に解釈すべきである。'
    )

    ms = RESULTS['model_selection']
    w_best = sum(1 for r in ms.values() if r['best_aic'] == 'Weibull')
    w_comp = sum(1 for r in ms.values() if r.get('w_daic') is not None and r['w_daic'] < 2)

    doc.add_heading('モデル選択結果', level=2)
    add_citation_paragraph(doc,
        f'体系的モデル選択の結果、AICによりワイブル分布は31政体中{w_best}政体で最良モデルとして'
        f'選択された。重要なことに、{w_comp}政体においてΔAIC < 2であり、データセットの大部分で'
        'ワイブルは統計的に競争力のあるフィットを提供した。ワイブルが明らかに不利な政体'
        '（ΔAIC > 4）は日本（天皇）と米国大統領であった。'
    )

    papacy = RESULTS['papacy']
    doc.add_heading('教皇庁：1000年以前vs以後', level=2)
    add_citation_paragraph(doc,
        f'1000年以前の教皇庁データ（N = {papacy["pre1000"]["n"]}）の分析結果、'
        f'k = {papacy["pre1000"]["k"]:.3f}（95% CI [{papacy["pre1000"]["ci"][0]:.2f}, '
        f'{papacy["pre1000"]["ci"][1]:.2f}]）であり、指数分布の基準値k = 1に近い値を示した。'
        f'これはコンクラーベ制度以前の教皇継承がほぼメモリレス（ランダム）であったことを示唆する。'
    )

    usa = RESULTS['usa_sensitivity']
    doc.add_heading('選挙制vs世襲制', level=2)
    add_citation_paragraph(doc,
        f'米国大統領を除外した選挙制政体のプールされた形状パラメータは'
        f'k = {usa["without"]:.3f}であり、世襲制のk = {usa["hereditary"]:.3f}とほぼ同等であった。'
    )

    ss = RESULTS['sample_size']
    doc.add_heading('標本サイズと信頼区間の評価', level=2)
    add_citation_paragraph(doc,
        f'31政体中{ss["n_crossing"]}政体の95%信頼区間がk = 1の閾値を跨いでおり、'
        'ポイント推定値のみでは減少ハザードと増加ハザードのレジームを信頼性をもって'
        '識別できないことが示された。'
    )

    # Discussion
    doc.add_heading('考察', level=1)
    add_citation_paragraph(doc,
        '本研究は、Salehの先駆的な単一政体研究を{1} 5大陸31政体に拡張した、統治者在位期間に'
        '対するワイブル信頼性分析の初の体系的グローバル応用である。'
    )
    add_citation_paragraph(doc,
        'モデル選択の結果は、ワイブル分布が普遍的に最良のフィットではないものの、大多数の'
        '政体で競争力のあるフィットを提供することを示した。ワイブルの解釈上の利点——形状'
        'パラメータkがハザード関数を介してバスタブ曲線フレームワークに直接対応すること——は、'
        '対数正規やガンマ分布が統計的に若干優れるフィットを提供する場合でも、政体間比較の'
        '主要な分析ツールとしての使用を正当化する。'
    )

    saleh = RESULTS['roman_saleh']
    add_citation_paragraph(doc,
        f'ローマ帝国に対する我々の単一ワイブルフィット（k = {saleh["single_k"]:.3f}）は、'
        f'Salehの混合ワイブルモデルと対照的である。{1} '
        '主要な方法論的差異は結果変数にある：Salehは暴力的死亡までの時間を分析したのに対し、'
        '我々は終了原因を問わない総在位期間を分析した。政体間比較という我々の目的のためには、'
        '簡潔な単一パラメータモデルが適切である。'
    )

    # Limitations
    doc.add_heading('限界', level=1)
    add_citation_paragraph(doc,
        '本研究にはいくつかの限界がある。第一に、歴史データの質と完全性は政体間で大きく異なる。'
        '第二に、継承メカニズムの二値分類は歴史的継承慣行の豊かな多様性を単純化しすぎている。'
        '第三に、一部の政体は定性的レジーム分類を支持するには標本サイズが不十分である。'
        '第四に、現代の議会制民主主義（西欧の首相など）がデータセットに含まれておらず、'
        '現代ガバナンスの非対称的カバレッジとなっている。首相の在任期間は根本的に異なる'
        '終了メカニズム（不信任決議、党首選挙）を伴うため、今後の研究で別途扱うべきである。'
    )

    # Conclusion
    doc.add_heading('結論', level=1)
    add_citation_paragraph(doc,
        'ワイブル分布は3千年にわたる31政体の政治的安定性を定量化するための堅牢で解釈可能な'
        '枠組みを提供することが実証された。体系的モデル選択により、大多数の政体でワイブル分布が'
        '競争力のあるフィットを提供し、形状パラメータkは暴力的継承の歴史的パターンを追跡する'
        '意味のある「安定性指標」を提供することが確認された。'
    )

    # Data availability
    doc.add_heading('データ利用可能性', level=1)
    add_citation_paragraph(doc,
        '本研究で使用した全在位期間データは公的に利用可能な歴史的資料から編纂された。'
        '完全なデータセットと分析コードは '
        'https://github.com/bougtoir/sovereign-weibull-analysis で利用可能である。'
    )

    # References (same as EN)
    doc.add_heading('参考文献', level=1)
    p = doc.add_paragraph()
    run = p.add_run('[英語版と同一の参考文献リスト]')
    run.font.size = Pt(10)
    run.italic = True

    # Ethical
    doc.add_heading('倫理的承認', level=1)
    p = doc.add_paragraph()
    run = p.add_run('本論文にはヒトを対象とした研究は含まれていない。')
    run.font.size = Pt(11)

    path = os.path.join(OUT_DIR, 'hassc_manuscript_ja_v4.docx')
    doc.save(path)
    print(f"Saved: {path}")
    return path


if __name__ == '__main__':
    generate_ja_manuscript()
